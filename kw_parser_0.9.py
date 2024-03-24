from flashtext import KeywordProcessor
from docx import Document
import re
import regex as reg
import os
import requests
from plyer import filechooser
from kivy.uix.boxlayout import BoxLayout
from kivy.core.window import Window
from kivy.metrics import dp
from kivymd.app import MDApp
from kivymd.uix.button import MDRaisedButton
from kivymd.uix.datatables import MDDataTable
from kivymd.uix.label import MDLabel
from kivymd.uix.textfield import MDTextField
import kivymd.icon_definitions
from kivymd.uix.dialog import MDDialog
from kivymd.uix.navigationrail import MDNavigationRail, MDNavigationRailItem
from kivymd.uix.scrollview import MDScrollView

"""
EXPORT WITH:
pyinstaller -F -w --hidden-import=plyer.platforms.win.filechooser --hidden-import=kivymd.uix.dropdownitem .\KW_Parser_0.8.py
"""

"""TODO:
-program causes memory leak on resize and reset, because fuck you 
"""
class SEO:
    info = []
    def __init__(self, meta, title_tag):
        self.meta = meta
        self.title_tag = title_tag

        if "%%currentyear%%" in meta:
            self.len_meta = len(meta) - 11
        else:
            self.len_meta = len(meta)

        if "%%currentyear%%" in title_tag:
            self.len_title_tag = len(title_tag) - 11
        else:
            self.len_title_tag = len(title_tag)

class Bolding_Errors:
    all_bolding_errors = []
    def __init__(self, line_string, bolding_error):
        self.line_string = line_string
        self.bolding_error = bolding_error
        
class Links:
    all_links = []
    document_ids = []
    domain = ""
    def __init__(self, a_tag):
        self.a_tag = a_tag
        self.href_content = self.extract_href(a_tag)
        self.anchor = self.extract_anchor(a_tag)
        self.status = self.check_status(self.href_content, self.domain)

    def __str__(self):
        return f'"{self.href_content}" with anchor "{self.anchor}" returns status: {self.status} \n'
    
    def extract_href(self, a_tag):
        href_pattern = r'<a\s+(?:[^>]*?\s+)?href="([^"]*)"'
        if match := re.findall(href_pattern, a_tag):
            return match[0]
        else:
            return "Error in href argument"
        
    def extract_anchor(self, a_tag):
        anchor_pattern = r'<a[^>]*>(.*?)<\/a>'
        if match := re.findall(anchor_pattern, a_tag):
            return match[0]
        else:
            return "Error in anchor text"
    
    def check_status(self, href_content, domain):
        if href_content.startswith("/"):
            try:
                url = domain + href_content
                response = requests.head(url, allow_redirects=True, timeout=5)
                if response.status_code == 200:
                    return f'[color=90EE90]Page exists[/color]'
                else:
                    return f"[color=ff0000]Page doesn't exist[/color]"
            except requests.RequestException as e:
                return f'Error checking {url}: {e}'
        elif href_content.startswith("#"):
            
            if href_content[1:] in self.document_ids:
                return "[color=90EE90]Internal link matches ID in content[/color]"
            else:
                return "[color=ff0000]Internal link does NOT MATCH ID in content[/color]"
        elif href_content.startswith("https://"):
            try:
                response = requests.head(href_content, allow_redirects=True, timeout=5)
                if response.status_code == 200:
                    return f'[color=90EE90]Page exists[/color]'
                else:
                    return f"[color=ff0000]Page doesn't exist[/color]"
            except requests.RequestException as e:
                return f'Error checking {href_content}: {e}'
        

class Paragraphs:
    long_paragraphs = []
    def __init__(self, paragraph: str, paragraph_len: int):
        self.paragraph = paragraph
        self.paragraph_len = paragraph_len

class Keyword:
    final_table = ""
    complete_data = []
    def __init__(self, keyword, exact_match=0, partial_match=0) -> None:
        self.keyword = keyword
        self.exact_match = exact_match
        self.partial_match = partial_match
        
    def __str__(self):
        return str([self.keyword, self.exact_match, self.partial_match])

class KeywordParser(MDApp): 
    def build(self):
        """
        A little explanation:
        the app has a main window (window0), within it are located the different sections
        left_side is the left half of the scren, where you enter files for parsing
        the right half (true_right_side) houses two additional sections, 
        one for the table visualisation (keyword_table) and for the buttons below it(right_side_buttons)
        """
        Window.size = (1400,900)
        self.window0 = BoxLayout()

        #entire left half is coded below
        self.left_side = BoxLayout(orientation='vertical',size_hint=(.25,1), spacing=10, padding = 10)        
        self.window0.add_widget(self.left_side) #this is the entire left section
        
        self.text_field = MDLabel(text="[color=000000]Parse keywords with my blessing", 
                                font_size=24, markup=True,pos_hint={'center_x':0.5,'center_y':0.2}, 
                                size_hint=(1,0.2), font_style='H3')
        self.left_side.add_widget(self.text_field) #title label

        self.left_side_input = BoxLayout(orientation='vertical', size_hint=(1, 0.7))
        self.left_side.add_widget(self.left_side_input) #element that houses all of the input elements, text fields and their buttons

        self.site = MDTextField(hint_text="Enter domain for article")
        self.left_side_input.add_widget(self.site)

        self.site_button = MDRaisedButton(text="Enter site", md_bg_color='red',
                                          pos_hint={'center_x':0.8, 'center_y': 0.8})
        self.site_button.bind(on_release=self.pass_site)
        self.left_side_input.add_widget(self.site_button)
        
        self.kw_list = MDTextField(hint_text="Enter file with keywords list", readonly=True)
        self.left_side_input.add_widget(self.kw_list) #text input for kw list file

        self.kw_selection_button = MDRaisedButton(text="Select KW file", md_bg_color='red',
                                                  pos_hint={'center_x':0.8,'center_y':0.8})
        self.kw_selection_button.bind(on_release=self.select_kw_list_file)
        self.left_side_input.add_widget(self.kw_selection_button) #button for selecting a kw list file
        
        self.docx = MDTextField(hint_text = "Enter .docx file to parse", readonly=True, size_hint=(1,.2))
        self.left_side_input.add_widget(self.docx) #text input for docx file

        self.docx_selection_button = MDRaisedButton(text="Select docx file",md_bg_color='red',
                                                    pos_hint={'center_x':0.8,'center_y':0.8})
        self.docx_selection_button.bind(on_release=self.select_docx_file)
        self.left_side_input.add_widget(self.docx_selection_button) #button for selecting docx file  

        self.parse_button = MDRaisedButton(text="Start KW parse", size_hint=(0.8,0.1), 
                                           md_bg_color='red',pos_hint={'center_x':0.5,'center_y':0.5})
        self.parse_button.bind(on_release=self.main_parse)
        self.left_side.add_widget(self.parse_button) #parse button

        
        #entire right half is coded below
        self.true_right_side = BoxLayout(size_hint=(.75,1))
        self.window0.add_widget(self.true_right_side) #this is the entire right section
        
        self.right_side = MDNavigationRail(
            MDNavigationRailItem(text="Home", icon="home", on_press=self.home_button),                                           
            MDNavigationRailItem(text="Keywords",icon="clipboard-text", on_press=self.show_table),
            MDNavigationRailItem(text="Boldings",icon="format-bold", on_press=self.show_bolding),
            MDNavigationRailItem(text="<p> length", icon="format-paragraph",on_press=self.show_paragraphs),
            MDNavigationRailItem(text="Links", icon="link-variant",on_press=self.show_links),
            MDNavigationRailItem(text="SEO", icon="google",on_press=self.show_seo)
            )
        self.true_right_side.add_widget(self.right_side) #not entirely sure what this is, houses the keyword_table section
        
        self.display_section = MDScrollView(scroll_y=1)
        self.true_right_side.add_widget(self.display_section)

        #below are all the elements that can be displayed in self.display_section
        self.keyword_table = MDDataTable(
            #use_pagination=True,
            rows_num=150,
            column_data = [
            ("Keyword", dp(80)),
            ("Exact Matches", dp(30)),
            ("Partial Matches", dp(30)),
            ("Total Matches", dp(30))
        ]
        ) 
        
        self.bolding_view = MDLabel(markup=True,padding = 30,size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.bolding_view.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1])) #don't touch, this makes the Label scrollable, don't know what it does specifically, it just works

        self.paragraph_view = MDLabel(markup=True, padding = 30,size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.paragraph_view.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))#don't touch, this makes the Label scrollable, don't know what it does specifically, it just works

        self.links_view = MDLabel(markup=True, padding = 30, size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.links_view.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))#don't touch, this makes the Label scrollable, don't know what it does specifically, it just works

        self.seo = MDLabel(markup=True, padding = 30, size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.seo.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))
   
        return self.window0
    
    def home_button(self, *args):
        self.display_section.clear_widgets()

    def get_seo(self, *args):
        if SEO.info:
            details = f''
            
            if SEO.info.len_title_tag > 60:
                description = f'\n\nTitle tag is [b][color=ff0000]{SEO.info.len_title_tag} characters. Target length is 60[/color][/b]. Try shortening. \n\n {SEO.info.title_tag} \n\n'
                details += description
            elif SEO.info.len_title_tag < 50:
                description = f'\n\nTitle tag is [b][color=ff0000]{SEO.info.len_title_tag} characters. Target length is 60[/color][/b]. Try lenghtening. \n\n {SEO.info.title_tag} \n\n'
                details += description
            else:
                description = f'\n\nTitle tag is within normal character range, [b][color=90EE90]currently {SEO.info.len_title_tag} characters[/color][/b]. \n\n {SEO.info.title_tag} \n\n'
                details += description
            if SEO.info.len_meta > 160:
                meta_desc = f'\n\nMeta description is [b][color=ff0000]{SEO.info.len_meta} characters. Target length is 155-160 max[/color][/b]. Try shortening. \n\n {SEO.info.meta} \n\n'
                details += meta_desc
            elif SEO.info.len_meta < 150:
                meta_desc = f'\n\nMeta description is [b][color=ff0000]{SEO.info.len_meta} characters. Target length is 155-160 max[/color][/b]. Try lenghtening. \n\n {SEO.info.meta} \n\n'
                details += meta_desc
            else:
                meta_desc = f'\n\nMeta description is within normal target range, [b][color=90EE90]currently {SEO.info.len_meta} characters[/color][/b]. \n\n {SEO.info.meta} \n\n'
                details += meta_desc
            self.seo.text = details

    def show_seo(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.seo)

    def get_links(self, *args):
        if Links.all_links:
            links = ""
            for i in Links.all_links:
                link_presentation = f'\n[b]{i.href_content}[/b] with anchor [b]{i.anchor}[/b] returns status: [b]{i.status}[/b]\n'
                links+=link_presentation
            self.links_view.text = links
    
    def show_links(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.links_view)

    def get_paragraphs(self, *args):
        if Paragraphs.long_paragraphs:
            too_long_paragraphs = ""
            for i in Paragraphs.long_paragraphs:
                error_message = f'[b] Paragraph too long. {i.paragraph_len} words with target 70[/b]: {i.paragraph}\n'
                too_long_paragraphs+=error_message
            self.paragraph_view.text=too_long_paragraphs

    def show_paragraphs(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.paragraph_view)
        
    def show_bolding(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.bolding_view)

    def show_table(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.keyword_table)
        
    def get_bolding(self, *args):
        if Bolding_Errors.all_bolding_errors:            
            error_messages = ''
            for i in Bolding_Errors.all_bolding_errors:                
                error = f'[b]{i.bolding_error}[/b]: {i.line_string}\n'
                error_messages+=error
            self.bolding_view.text = error_messages            
        else:
            self.bolding_view.text = "No bolding errors found"            
            
    def pass_site(self, *args):
        link = self.site.text.strip()
        if link[0:8] != 'https://':
            link = 'https://' + link
        Links.domain = link

    #below is file selection fuction
    def select_kw_list_file(self, *args):
        #from plyer import filechooser
        file = filechooser.open_file(on_selection=self.selected_kw_list, filters=['*.txt'])
        if file == []:
            pass
        else:
            self.kw_list.text = file[0]
    
    #below is file selection fuction
    def selected_kw_list(self, selection):
        return selection
    
    #below is file selection fuction
    def select_docx_file(self, *args):

        file = filechooser.open_file(on_selection=self.selected_docx_file, filters=['*.docx'])
        if file == []:
            pass
        else:
            self.docx.text=file[0]

    #below is file selection fuction
    def selected_docx_file(self, selection):
        return selection

    def get_final_table(self, *args): #this function adds the parsed data to the KW table
        
        for i in Keyword.complete_data:
            self.keyword_table.add_row(
                (i.keyword, i.exact_match, i.partial_match, i.exact_match + i.partial_match)
            )
        
        #self.table_button.md_bg_color='black'
        self.parse_button.md_bg_color='red'
    
    def main_parse(self, *args): #main function, runs required functions in order
        try:            
            keyword_list = self.get_keywords_from_list()
            KW = self.pass_keywords_to_class(keyword_list)
            new_doc = self.word_document()
            formatted_doc = self.format(new_doc)
            self.exact_match_parse(formatted_doc, KW)
            self.partial_match_parse(formatted_doc, keyword_list, KW)
            Keyword.complete_data = KW
            del KW            
            os.remove('formatted_temp.txt')
            os.remove('temp.txt')
            self.get_final_table()
            self.get_bolding()
            self.get_paragraphs()
            self.get_links()
            self.get_seo()
            
        except FileNotFoundError:
            self.file_not_found_error = MDDialog(text="File not found")
            self.file_not_found_error.open()
        
    def get_keywords_from_list(self): #txt file which passed and every line is extracted as a keyword
        list_file = self.kw_list.text
        keyword_list = []
        with open(list_file) as file:
            file = file.readlines()
            for line in file:
                keyword_list.append(line.strip())
        return keyword_list


    def pass_keywords_to_class(self, keyword_list): #creates Keyword class instances and the variable that houses them
        KW = []
        for keyword in keyword_list:
            KW.append(Keyword(keyword))
        return KW

    def word_document(self):
        new_doc = open('temp.txt', 'w', errors='ignore')
        document = Document(self.docx.text)

        meta = ""
        title_tag = ""
        for number, table in enumerate(document.tables): #this fucking thing parses the tables for SEO title and meta description, shit on efficiency
            for num, row in enumerate(table.rows):
                for n, cell in enumerate(row.cells):
                    if cell.text.lower() == "meta:":
                        meta = document.tables[number].rows[num].cells[n+1].text
                    if cell.text.lower() == "title tag:":
                        title_tag = document.tables[number].rows[num].cells[n+1].text
        SEO_info = SEO(meta, title_tag)
        SEO.info = SEO_info
        del meta, title_tag, SEO_info

        for i in document.paragraphs:
            new_doc.write(i.text + '\n')
        return 'temp.txt'
    
    def format(self, new_doc): #checks for bolding errors, then 
        with open(new_doc, 'r') as original_file, open('formatted_temp.txt', 'w') as formatted_file:
            long_paragraphs = []
            bolding_errors=[]
            links = []
            class_links = []
            
            for line in original_file.readlines():
                if re.findall(r'id="([^"]+)"', line):
                    Links.document_ids.extend(re.findall(r'id="([^"]+)"', line))
                if re.findall(r'<a[^>]*>.*?</a>',line): #this find all <a> tags in the content
                    links.extend(re.findall(r'<a[^>]*>.*?</a>', line))
                if len(line.split(" ")) > 70: #paragraph length check starts here
                    long_paragraphs.append(Paragraphs(line, len(line.split(" "))))
                if line == "\n" or line.startswith("<h") or line.startswith("<li"): #bolding check starts here
                    continue
                elif len(line.split(" "))<25:
                    continue
                elif line.count("<strong>") != line.count("</strong>"):
                    bolding_errors.append(Bolding_Errors(line, "Unmatched <strong> tags"))
                else:
                    pattern = r'<strong>.*?</strong>'
                    found_patterns = reg.findall(pattern, line, overlapped=True)
                    if len(found_patterns)<1:
                        bolding_errors.append(Bolding_Errors(line, "No boldings"))
                    elif len(found_patterns)>1:
                        bolding_errors.append(Bolding_Errors(line, "Too many boldings"))
                    else:
                        pass #bolding check ends here
                if line == "":
                    continue
                clean = re.compile('<.*?>')
                line = re.sub(clean,'', line)
                line = re.sub("[,.?:']",'', line)
                line = re.sub("\[[0-9]\]", '', line)
                formatted_file.write(line)
                        
            for link in links:
                class_links.append(Links(link))
            Links.all_links = class_links
            Bolding_Errors.all_bolding_errors = bolding_errors
            Paragraphs.long_paragraphs = long_paragraphs
            
            del long_paragraphs, bolding_errors, links, class_links

        return 'formatted_temp.txt'    

    def exact_match_parse(self, formatted_file, KW): #
        
        formatted_document = open(formatted_file).read()
        keyword_processor = KeywordProcessor()
        for i in KW:
            keyword_processor.add_keyword(i.keyword)
        kw_found = keyword_processor.extract_keywords(formatted_document)
        for i in KW:
            i.exact_match += kw_found.count(i.keyword)

    def partial_match_parse(self, formatted_file, keyword_list, KW):
        kw_patterns = []
        for keyword in keyword_list:
            separated_kw = keyword.split(" ")
            kw_patterns.append('(?: |(?:(?: \w* )(?:\w* )?))'.join(separated_kw))
        with open(formatted_file, 'r') as file:
            file = file.read()
            for number, pattern in enumerate(kw_patterns):
                matches = re.findall(pattern, file, re.IGNORECASE)
                matches = [x.lower() for x in matches]                
                for match in matches:
                    if match in keyword_list:
                        continue
                    else:
                        KW[number].partial_match +=1

if __name__ == "__main__":
    KeywordParser().run()
