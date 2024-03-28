from flashtext import KeywordProcessor
from docx import Document
import re
import regex as reg
import os
import requests
from plyer import filechooser 
import difflib
#import plyer.platforms.macosx.filechooser
from kivy.uix.boxlayout import BoxLayout
from kivy.core.window import Window
from kivy.metrics import dp
from kivymd.app import MDApp
from kivymd.uix.button import MDRaisedButton
from kivymd.uix.datatables import MDDataTable
from kivymd.uix.label import MDLabel
from kivymd.uix.textfield import MDTextField
import kivymd.icon_definitions
from kivymd.uix.dialog import MDDialog
from kivymd.uix.navigationrail import MDNavigationRail, MDNavigationRailItem
from kivymd.uix.scrollview import MDScrollView
from kivymd.uix.list import MDList, OneLineListItem
from kivymd.uix.boxlayout import MDBoxLayout
#import pdb; pdb.set_trace()
"""
EXPORT TO EXE WITH:
pyinstaller -F -w --hidden-import=plyer.platforms.win.filechooser --hidden-import=kivymd.uix.dropdownitem .\KW_Parser_0.8.py
"""

"""TODO:
-program causes memory leak on resize and reset, because fuck you
"""



global database

database = ["Admiral", 
"All British Casino", 
"AHTI Games", 
"Aspers Casino", 
"All Slots", 
"BC.GAME", 
"Bao Casino", 
"Betsafe", 
"BetChan", 
"Bethard", 
"Betsson", 
"Betnero", 
"Betway", 
"BitStarz", 
"Bingoal", 
"Billion Casino", 
"Borgata Casino", 
"Boomerang", 
"bwin", 
"BetVictor", 
"BetNation", 
"Betano", 
"bCasino", 
"bet365 ", 
"BetRivers Casino", 
"BoyleSports Casino", 
"Captain Cooks Casino", 
"Casino 2020", 
"Casino Action", 
"Casino Classic", 
"Casino Cruise", 
"Casino Gods", 
"Casino Heroes", 
"Casino Joy", 
"Casino Lab", 
"CasinoLuck", 
"Casino Of Dreams", 
"Casino GranMadrid ", 
"CasinoClub", 
"Casino Midas", 
"Casiplay", 
"Caesars Casino", 
"Crystal Slots", 
"Casino Planet", 
"cky", 
"Conquer Casino", 
"CobraSpins ", 
"Codere", 
"Cloud Casino", 
"Clover Casino", 
"Casitsu", 
"Casombie ", 
"ComeOn!", 
"Casoo ", 
"Casino.com", 
"Casimba", 
"Casumo", 
"Circus", 
"Casoola", 
"Dazzle Casino", 
"DrückGlück", 
"Dunder", 
"DraftKings Casino", 
"Dream Vegas", 
"Duxcasino", 
"Euro Palace", 
"Empire Casino", 
"EnergyCasino", 
"efbet", 
"Euslot", 
"Frank & Fred", 
"Frank Casino", 
"Fortune Mobile Casino", 
"Fair Play", 
"FoggyStar", 
"Fortuna", 
"Fun Casino", 
"Ganabet", 
"Gaming Club", 
"Glimmer Casino", 
"Goliathcasino", 
"Grand Hotel Casino", 
"GetSlots", 
"GG.BET", 
"Gioco Digitale", 
"Goldrun Casino", 
"GoDaddy", 
"GoSlotty", 
"Genesis Casino", 
"Gala Casino", 
"Grosvenor", 
"Griffon Casino", 
"Golden Nugget", 
"Holland Casino", 
"Hollywood casino", 
"Harrah's ", 
"Hard Rock Online Casino", 
"Heart of Casino", 
"Hopa", 
"Hyper Casino", 
"ICE Casino", 
"InterCasino", 
"Jack's Casino", 
"Jackpot Mobile Casino", 
"Jambo Casino", 
"Joo Casino", 
"Kansino", 
"Kaiser Slots", 
"King Casino", 
"King Billy", 
"Karamba", 
"lapalingo ", 
"Luxury Casino", 
"LV BET", 
"LuckyDays", 
"Locowin", 
"LeoVegas", 
"Ladbrokes", 
"LiveScore Bet ", 
"LuckyNiki", 
"Mega Casino", 
"MEGAWAYS Casino", 
"Mummys Gold Casino", 
"MONOPOLY Casino", 
"Monster Casino", 
"MansionCasino", 
"MagicJackpot ", 
"MagicRed", 
"Megaslot", 
"MegaRush", 
"MERKUR SLOTS ", 
"MobileBet", 
"Mr Jack Vegas", 
"Mr Bet", 
"MaxBet", 
"Mr Bit ", 
"mr.play", 
"Mr Green", 
"N1 Casino", 
"NextCasino", 
"NineCasino", 
"Nitro Casino", 
"NightRush", 
"NetBet", 
"NYspins Casino", 
"One Casino", 
"Ocean Online Casino", 
"Pala", 
"Paddy Power", 
"Paf Casino", 
"PinoCasino ", 
"PlayAmo", 
"PlayZilla", 
"Prank Casino", 
"PublicWin", 
"Plush Casino", 
"PartyCasino", 
"Playzee", 
"Play Jango", 
"Platinum Play", 
"PlayOJO", 
"Queen Vegas", 
"Quatro Casino", 
"Queenplay ", 
"Regent Play", 
"River Belle Casino", 
"Rise Casino", 
"Rainbow Casino", 
"Rolling Slots", 
"Royal Panda", 
"Ruby Fortune", 
"SuperCasino", 
"Stardust Casino", 
"Swift Casino", 
"Scatters", 
"Simba Games", 
"Sisal", 
"SlotV Casino", 
"Slotty Vegas", 
"Stars Casino", 
"Sons Of Slots", 
"Spin Rider", 
"Spin Samurai ", 
"Sportingbet", 
"StarCasinò", 
"StarVegas", 
"Sugar Casino", 
"Sky Vegas", 
"Sunmaker", 
"Sunnyplayer", 
"Surf Casino", 
"SpinYoo", 
"SlotsMillion", 
"TOTO Casino", 
"tombola", 
"The Casino MK", 
"The Grand Ivy", 
"TradaCasino", 
"UK Casino Club", 
"Unibet", 
"Virgin Casino", 
"Vegadream", 
"Vulkan Vegas", 
"Villento Casino", 
"Videoslots", 
"Voodoo Dreams", 
"Vlad Cazino", 
"WynnBET Casino", 
"Wanabet", 
"Wazamba", 
"WestCasino", 
"Wunderino", 
"Wild Fortune ", 
"Winmasters", 
"Winbet", 
"WildTornado", 
"William Hill", 
"Win British", 
"Yako Casino", 
"YaaCasino", 
"Yeti Casino", 
"Yukon Gold Casino", 
"Zodiac Casino", 
"ZodiacBet ", 
"ZEbet", 
"22BET", 
"20Bet", 
"1xBet", 
"1Bet", 
"14Red", 
"123 Spins", 
"666 Casino", 
"777 Casino", 
"1xSlots", 
"32Red", 
"21.com", 
"711", 
"10bet", 
"888casino", 
"Apple Pay", 
"AstroPay", 
"American Express", 
"Bitcoin", 
"Boku", 
"CashtoCode", 
"Citadel", 
"CashPay", 
"Payz ", 
"Google Pay", 
"Interac", 
"Jeton", 
"Klarna", 
"Inpay", 
"MiFinity", 
"Mastercard", 
"Maestro", 
"MuchBetter", 
"Neteller", 
"Neosurf", 
"paysafecard", 
"Paytm", 
"Siru Mobile", 
"SMS", 
"Skrill", 
"TopPay", 
"Trustly", 
"UPayCard", 
"Visa", 
"WebMoney", 
"Zimpler", 
"Pay N Play", 
"Pay By SMS", 
"e-Wallet", 
"Ash Gaming", 
"Adoptit Publishing", 
"Authentic Gaming", 
"AMATIC", 
"Ainsworth Game Technology", 
"Barcrest", 
"BetGames", 
"Belatra Games", 
"Booming Games", 
"Bally", 
"Betsoft", 
"BF Games", 
"BGaming", 
"Bla Bla Bla Studios", 
"Booongo", 
"Blueprint Gaming", 
"Big Time Gaming", 
"Crazy Tooth Studio", 
"Core Gaming", 
"ELK Studios", 
"Endorphina", 
"Evoplay", 
"Ezugi", 
"Evolution", 
"Electric Elephant", 
"Fantasma Games", 
"Foxium", 
"Felt", 
"Gamevy", 
"GameArt", 
"Gamesys", 
"Genesis Gaming", 
"G GAMING", 
"Dragonfish", 
"Hacksaw Gaming", 
"Habanero", 
"Intouch Games", 
"Iron Dog Studio", 
"iSoftBet", 
"Igrosoft", 
"IGT", 
"Just For The Win", 
"Kalamba Games", 
"KONAMI", 
"Lightning Box", 
"Leap", 
"Leander Games", 
"Mascot Gaming", 
"Microgaming", 
"MrSlotty", 
"Nolimit City", 
"NetEnt", 
"Nektan", 
"Neogames", 
"Netgame", 
"Northern Lights Gaming", 
"Novoline", 
"NextGen Gaming", 
"Novomatic", 
"Old Skool Studios", 
"Pariplay", 
"Platipus", 
"Probability", 
"Push Gaming", 
"Pragmatic Play", 
"Playson", 
"Platipus Gaming", 
"Playtech", 
"Play'n GO", 
"Quickspin", 
"Quickfire", 
"Rabcat", 
"Random Logic", 
"Relax Gaming", 
"Red Tiger", 
"Realistic Games", 
"ReelPlay", 
"Yggdrasil", 
"Sigma Games", 
"Skillzzgaming", 
"Spinomenal", 
"Slotvision", 
"STHLMGAMING", 
"TonyBet", 
"Thunderkick", 
"Wazdan", 
"Wild Streak Gaming", 
"BeGambleAware", 
"Cloudflare", 
"Comodo", 
"Curaçao eGaming", 
"DigiCert", 
"eCOGRA ", 
"Gamblers Anonymous", 
"Gambling Therapy", 
"GamCare", 
"GAMSTOP", 
"GoGetSSL", 
"HydrantID", 
"iTech Labs", 
"MGA", 
"RNG Certified", 
"Sectigo", 
]


class SEO:
    info = []
    def __init__(self, meta, title_tag):
        self.meta = meta
        self.title_tag = title_tag

        if "%%currentyear%%" in meta:
            self.len_meta = len(meta) - 11
        else:
            self.len_meta = len(meta)

        if "%%currentyear%%" in title_tag:
            self.len_title_tag = len(title_tag) - 11
        else:
            self.len_title_tag = len(title_tag)

class Bolding_Errors:
    all_bolding_errors = []
    def __init__(self, line_string, bolding_error):
        self.line_string = line_string
        self.bolding_error = bolding_error

class Spelling:
    spelling_errors = []
    def __init__(self, caught_phrase, suggested_phrase, paragraph):
        self.caught_phrase = caught_phrase
        self.suggested_phrase = suggested_phrase
        self.paragraph = paragraph

class Links:
    link_errors = []
    all_links = []
    document_ids = []
    domain = ""
    slug = ""
    def __init__(self, a_tag):
        self.a_tag = a_tag
        self.href_content = self.extract_href(a_tag)
        self.anchor = self.extract_anchor(a_tag)
        self.status = self.check_status(self.href_content, self.domain)

    def __str__(self):
        return f'"{self.href_content}" with anchor "{self.anchor}" returns status: {self.status} \n'

    def extract_href(self, a_tag):
        href_pattern = r'<a\s+(?:[^>]*?\s+)?href="([^"]*)"'
        if match := re.findall(href_pattern, a_tag):
            return match[0]
        else:
            return "Error in href argument"

    def extract_anchor(self, a_tag):
        anchor_pattern = r'<a[^>]*>(.*?)<\/a>'
        if match := re.findall(anchor_pattern, a_tag):
            return match[0]
        else:
            return "Error in anchor text"

    def check_status(self, href_content, domain):
        if href_content.startswith("/"):
            try:
                url = domain + href_content
                response = requests.head(url, allow_redirects=True, timeout=5)
                if response.status_code == 200:
                    return f'[color=90EE90]Page exists[/color]'
                else:
                    return f"[color=ff0000]Page doesn't exist[/color]"
            except requests.RequestException as e:
                return f'Error checking {domain}: {e}'
        elif href_content.startswith("#"):
            if href_content[1:] in self.document_ids:
                return "[color=90EE90]Internal link matches ID in content[/color]"
            else:
                return "[color=ff0000]Internal link does NOT MATCH ID in content[/color]"
        elif href_content.startswith("https://"):
            try:
                response = requests.head(href_content, allow_redirects=True, timeout=5)
                if response.status_code == 200:
                    return f'[color=90EE90]Page exists[/color]'
                else:
                    return f"[color=ff0000]Page doesn't exist[/color]"
            except requests.RequestException as e:
                return f'Error checking {href_content}: {e}'
        for jump in self.document_ids:
            if self.document_ids.count(jump) > 1:
                self.link_errors.append(f'\n[b][color=ff0000]#{id}[/color] is assigned more than once[/b]. Check text.\n')
            else:
                continue


class Paragraphs:
    long_paragraphs = []
    h_tag_errors = []
    country_modifier_errors = []
    country = []
    countries = [['Argentina', 'AR'],
 ['Armenia', 'AM'],
 ['Australia', 'AU'],
 ['Austria', 'AT'],
 ['Azerbaijan', 'AZ'],
 ['Belgium', 'BE'],
 ['Bosnia & Herzegovina', 'BA'],
 ['Brazil', 'BR'],
 ['Bulgaria', 'BG'],
 ['Canada', 'CA'],
 ['Chile', 'CL'],
 ['Colombia', 'CO'],
 ['Croatia', 'HR'],
 ['Cyprus', 'CY'],
 ['Czech Republic', 'CZ'],
 ['Denmark', 'DK'],
 ['Dominican Republic', 'DO'],
 ['Egypt', 'EG'],
 ['Estonia', 'EE'],
 ['Ethiopia', 'ET'],
 ['France', 'FR'],
 ['Germany', 'DE'],
 ['Ghana', 'GH'],
 ['Gibraltar', 'GI'],
 ['Greece', 'GR'],
 ['Hong Kong', 'HK'],
 ['Hungary', 'HU'],
 ['Ireland', 'IE'],
 ['Isle of Man', 'IM'],
 ['Italy', 'IT'],
 ['Jamaica', 'JM'],
 ['Japan', 'JP'],
 ['Kenya', 'KE'],
 ['Latvia', 'LV'],
 ['Liechtenstein', 'LI'],
 ['Lithuania', 'LT'],
 ['Luxembourg', 'LU'],
 ['Malta', 'MT'],
 ['Mexico', 'MX'],
 ['Moldova', 'MD'],
 ['Montenegro', 'ME'],
 ['Netherlands', 'NL'],
 ['Nigeria', 'NG'],
 ['Norway', 'NO'],
 ['Panama', 'PA'],
 ['Philippines', 'PH'],
 ['Poland', 'PL'],
 ['Portugal', 'PT'],
 ['Romania', 'RO'],
 ['Russia', 'RU'],
 ['Serbia', 'RS'],
 ['Singapore', 'SG'],
 ['Slovakia', 'SK'],
 ['Slovenia', 'SI'],
 ['South Africa', 'ZA'],
 ['Spain', 'ES'],
 ['Sweden', 'SE'],
 ['Switzerland', 'CH'],
 ['Tanzania', 'TZ'],
 ['Trinidad and Tobago', 'TT'],
 ['Uganda', 'UG'],
 ['Ukraine', 'UA'],
 ['United States', 'US','USA'],
 ['Zambia', 'ZM'],
 ['Zimbabwe', 'ZW'],
 ['Alabama', 'AL'],
 ['Alaska', 'AK'],
 ['Arizona', 'AZ'],
 ['Arkansas', 'AR'],
 ['California', 'CA'],
 ['Colorado', 'CO'],
 ['Connecticut', 'CT'],
 ['Delaware', 'DE'],
 ['Florida', 'FL'],
 ['Georgia', 'GA'],
 ['Hawaii', 'HI'],
 ['Idaho', 'ID'],
 ['Illinois', 'IL'],
 ['Indiana', 'IN'],
 ['Iowa', 'IA'],
 ['Kansas', 'KS'],
 ['Kentucky', 'KY'],
 ['Louisiana', 'LA'],
 ['Maine', 'ME'],
 ['Maryland', 'MD'],
 ['Massachusetts', 'MA'],
 ['Michigan', 'MI'],
 ['Minnesota', 'MN'],
 ['Mississippi', 'MS'],
 ['Missouri', 'MO'],
 ['Montana', 'MT'],
 ['Nebraska', 'NE'],
 ['Nevada', 'NV'],
 ['New Hampshire', 'NH'],
 ['New Jersey', 'NJ'],
 ['New Mexico', 'NM'],
 ['New York', 'NY'],
 ['North Carolina', 'NC'],
 ['North Dakota', 'ND'],
 ['Ohio', 'OH'],
 ['Oklahoma', 'OK'],
 ['Oregon', 'OR'],
 ['Pennsylvania', 'PA'],
 ['Rhode Island', 'RI'],
 ['South Carolina', 'SC'],
 ['South Dakota', 'SD'],
 ['Tennessee', 'TN'],
 ['Texas', 'TX'],
 ['Utah', 'UT'],
 ['Vermont', 'VT'],
 ['Virginia', 'VA'],
 ['Washington', 'WA'],
 ['West Virginia', 'WV'],
 ['Wisconsin', 'WI'],
 ['Wyoming', 'WY']]
    
    def __init__(self, paragraph: str, paragraph_len: int):
        self.paragraph = paragraph
        self.paragraph_len = paragraph_len

class Keyword:
    final_table = ""
    complete_data = []
    def __init__(self, keyword, exact_match=0, partial_match=0) -> None:
        self.keyword = keyword
        self.exact_match = exact_match
        self.partial_match = partial_match

    def __str__(self):
        return str([self.keyword, self.exact_match, self.partial_match])

class KeywordParser(MDApp):
    def build(self):
        """
        A little explanation:
        the app has a main window (window0), within it are located the different sections
        left_side is the LEFT HALF of the scren, where you enter files for parsing

        the RIGHT HALF (true_right_side) houses a navigation rail, which can use to switch between the different tabs.
        Most tabs in the right half are displayed using an MDScrollView element, except the spelling table, which can go fuck itself.
        """
        Window.size = (1400,900)
        self.window0 = BoxLayout()

        #entire left half is coded below
        self.left_side = BoxLayout(orientation='vertical',size_hint=(.25,1), spacing=10, padding = 10)
        self.window0.add_widget(self.left_side) #this is the entire left section

        self.text_field = MDLabel(text="[color=000000]Parse keywords with my blessing",
                                font_size=24, markup=True,pos_hint={'center_x':0.5,'center_y':0.2},
                                size_hint=(1,0.2), font_style='H3')
        self.left_side.add_widget(self.text_field) #title label

        self.left_side_input = BoxLayout(orientation='vertical', size_hint=(1, 0.7))
        self.left_side.add_widget(self.left_side_input) #element that houses all of the input elements, text fields and their buttons

        self.site = MDTextField(hint_text="Enter domain for article")
        self.left_side_input.add_widget(self.site)

        self.site_button = MDRaisedButton(text="Enter site", md_bg_color='red',
                                          pos_hint={'center_x':0.8, 'center_y': 0.8})
        self.site_button.bind(on_release=self.pass_site)
        self.left_side_input.add_widget(self.site_button)

        self.kw_list = MDTextField(hint_text="Enter file with keywords list", readonly=True)
        self.left_side_input.add_widget(self.kw_list) #text input for kw list file

        self.kw_selection_button = MDRaisedButton(text="Select KW file", md_bg_color='red',
                                                  pos_hint={'center_x':0.8,'center_y':0.8})
        self.kw_selection_button.bind(on_release=self.select_kw_list_file)
        self.left_side_input.add_widget(self.kw_selection_button) #button for selecting a kw list file

        self.docx = MDTextField(hint_text = "Enter .docx file to parse", readonly=True, size_hint=(1,.2))
        self.left_side_input.add_widget(self.docx) #text input for docx file

        self.docx_selection_button = MDRaisedButton(text="Select docx file",md_bg_color='red',
                                                    pos_hint={'center_x':0.8,'center_y':0.8})
        self.docx_selection_button.bind(on_release=self.select_docx_file)
        self.left_side_input.add_widget(self.docx_selection_button) #button for selecting docx file

        self.parse_button = MDRaisedButton(text="Start KW parse", size_hint=(0.8,0.1),
                                           md_bg_color='red',pos_hint={'center_x':0.5,'center_y':0.5})
        self.parse_button.bind(on_press=self.parse_button_color_change)
        self.parse_button.bind(on_release=self.main_parse)
        self.left_side.add_widget(self.parse_button) #parse button


        #entire right half is coded below
        self.true_right_side = BoxLayout(size_hint=(.75,1))
        self.window0.add_widget(self.true_right_side) #this is the entire right section

        self.right_side = MDNavigationRail(
            MDNavigationRailItem(text="Home", icon="home", on_press=self.home_button),
            MDNavigationRailItem(text="Keywords",icon="clipboard-text", on_press=self.show_table),
            MDNavigationRailItem(text="Boldings",icon="format-bold", on_press=self.show_bolding),
            MDNavigationRailItem(text="<p> length", icon="format-paragraph",on_press=self.show_paragraphs),
            MDNavigationRailItem(text="Links", icon="link-variant",on_press=self.show_links),
            MDNavigationRailItem(text="SEO", icon="google",on_press=self.show_seo),
            MDNavigationRailItem(text="Spelling", icon="spellcheck",on_press=self.show_spelling),
            MDNavigationRailItem(text="RESET", icon="close-box",on_press=self.program_reset)

            )
        self.true_right_side.add_widget(self.right_side) #not entirely sure what this is, houses the keyword_table section

        self.display_section = MDScrollView(scroll_y=1)
        self.true_right_side.add_widget(self.display_section)

        #below are all the elements that can be displayed in self.display_section
        self.keyword_table = MDDataTable(
            #use_pagination=True,
            rows_num=150,
            column_data = [
            ("Keyword", dp(80)),
            ("Exact Matches", dp(30)),
            ("Partial Matches", dp(30)),
            ("Total Matches", dp(30))
        ]
        )

        self.bolding_view = MDLabel(markup=True,padding = 30,size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.bolding_view.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1])) #don't touch, this makes the Label scrollable, don't know what it does specifically, it just works

        self.paragraph_view = MDLabel(markup=True, padding = 30,size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.paragraph_view.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))#don't touch, this makes the Label scrollable, don't know what it does specifically, it just works

        self.links_view = MDLabel(markup=True, padding = 30, size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.links_view.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))#don't touch, this makes the Label scrollable, don't know what it does specifically, it just works

        self.seo = MDLabel(markup=True, padding = 30, size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.seo.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))
        
        #below are the spelling tab elements
        self.spelling_main_container = BoxLayout() #this is the main container
        #this is the left side of the Spelling tab, which includes the mistakes and the paragraphs in which they are caught
        self.spelling = MDLabel(markup=True, padding = 30, size_hint_y=None, size_hint_x = 0.9, valign='top')
        self.spelling.bind(texture_size=lambda instance, value: setattr(instance, 'height', value[1]))

        self.spelling_container = MDScrollView()
        self.spelling_container.add_widget(self.spelling)
        self.spelling_main_container.add_widget(self.spelling_container)


        #this is the right side of the Spelling tab, with the list of correctly-spelled phrases
        self.spelling_right_side = MDBoxLayout(orientation='vertical')
        self.spelling_references_search_box = MDTextField(hint_text="Enter phrase and press ENTER", size_hint=(0.9,0.1), on_text_validate=self.on_search_enter)
        self.spelling_references_search_box.bind() 
        self.spelling_references = MDList()
        self.spelling_references_container = MDScrollView(size_hint=(0.95, 0.9))
        self.spelling_references_container.add_widget(self.spelling_references)

        self.spelling_right_side.add_widget(self.spelling_references_search_box)
        self.spelling_right_side.add_widget(self.spelling_references_container)
        self.spelling_main_container.add_widget(self.spelling_right_side)

        return self.window0

    #home button, clears widgets from tab, does nothing else
    def home_button(self, *args):
        self.display_section.clear_widgets()

    def parse_button_color_change(self, *args):
        self.parse_button.md_bg_color='black'

    #resets the window content
    def program_reset(self, *args): 
        self.root.clear_widgets()
        self.stop()
        self.__init__()
        self.run()

    #this checks html id against the slug to find repeating words
    def check_slug_for_id(self, id):
        slug = (re.sub("[-/]", " ", Links.slug)).split(" ")
        id = id.strip().split("-")
        for id_part in id:
            if id_part in slug:
                return True
            else:
                return False

    #this bundle of joy below is for the list element showing the commonly mispelled words
    #includes search function, data getter, and data display, last one is bound to a button
    def on_search_enter(self, *args):
        search_input = self.spelling_references_search_box.text
        self.filter_list(search_input)

    def filter_list(self, value, *args):
        # Filter the list based on the search field's text
        filtered_data = [item for item in database if value.lower() in item.lower()]
        self.populate_list(filtered_data)

    def populate_list(self, data=None):
        # Populate the list with items, optionally using filtered data
        self.spelling_references.clear_widgets()  # Clear existing list items
        if not data:
            data = database  # If no data is provided, use the full list
        for item in data:
            self.spelling_references.add_widget(OneLineListItem(text=item))

    def get_spelling_references(self, *args):
        for phrase in sorted(database):
            self.spelling_references.add_widget(
                OneLineListItem(text=f"{phrase}")
            )

    def get_spelling_errors(self, *args):
        if Spelling.spelling_errors:
            error_text = "[size=36][b]This method is imperfect, use with caution and your own judgement.\nReference potentially mispelled names against the database list[/b][/size] \n \n"

            for item in Spelling.spelling_errors:
                description = f'[b][color=ff0000]!!![/color][/b] Potential spelling or formatting mistake [b]|{item.caught_phrase}|[/b]. Found in: \n{item.paragraph} \n\n'
                error_text +=description
            self.spelling.text = error_text

    def show_spelling(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.spelling_main_container)
    
    

    #this gets the data for the SEO errors tab
    def get_seo(self, *args):
        if SEO.info:
            details = f''
            if '[year]' in SEO.info.title_tag or '[month]' in SEO.info.title_tag:
                description = f"\n\n[b][color=ff0000]Wrong date code[/color][/b] in meta description. \n\n{SEO.info.title_tag}\n\n"
            else:
                if SEO.info.len_title_tag > 60:
                    description = f'\n\nTitle tag is [b][color=ff0000]{SEO.info.len_title_tag} characters. Target length is 60[/color][/b]. Try shortening. \n\n {SEO.info.title_tag} \n\n'
                    details += description
                elif SEO.info.len_title_tag < 50:
                    description = f'\n\nTitle tag is [b][color=ff0000]{SEO.info.len_title_tag} characters. Target length is 60[/color][/b]. Try lenghtening. \n\n {SEO.info.title_tag} \n\n'
                    details += description
                else:
                    description = f'\n\nTitle tag is within normal character range, [b][color=90EE90]currently {SEO.info.len_title_tag} characters[/color][/b]. \n\n {SEO.info.title_tag} \n\n'
                    details += description

            if '[year]' in SEO.info.meta or '[month]' in SEO.info.meta:
                meta_desc = f"\n\n[b][color=ff0000]Wrong date code[/color][/b] in meta description. \n\n{SEO.info.meta}\n\n"
                details += meta_desc
            else:
                if SEO.info.len_meta > 160:
                    meta_desc = f'\n\nMeta description is [b][color=ff0000]{SEO.info.len_meta} characters. Target length is 155-160 max[/color][/b]. Try shortening. \n\n {SEO.info.meta} \n\n'
                    details += meta_desc
                elif SEO.info.len_meta < 150:
                    meta_desc = f'\n\nMeta description is [b][color=ff0000]{SEO.info.len_meta} characters. Target length is 155-160 max[/color][/b]. Try lenghtening. \n\n {SEO.info.meta} \n\n'
                    details += meta_desc
                else:
                    meta_desc = f'\n\nMeta description is within normal target range, [b][color=90EE90]currently {SEO.info.len_meta} characters[/color][/b]. \n\n {SEO.info.meta} \n\n'
                    details += meta_desc
            self.seo.text = details

    #this shows mistakes in SEO elemens
    def show_seo(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.seo)

    #this gets the text for the links errors that will be displayed in the tab
    def get_links(self, *args):
        links = ""
        if Links.link_errors:
            for paragraph in Links.link_errors:
                error = f'\n[b][color=ff0000]LINK ERROR. Incorrect <a> tag or multiple links[/color][/b] in paragraph: {paragraph}\n'
                links+=error
        
        if Links.all_links:    
            for i in Links.all_links:
                link_presentation = f'\n[b]{i.href_content}[/b] with anchor [b]{i.anchor}[/b] returns status: [b]{i.status}[/b]\n'
                links+=link_presentation
        
        if Links.document_ids:
            for id in Links.document_ids:
                if self.check_slug_for_id(id) == True:
                    error = f'\n[b][color=ff0000]!!![/color][/b] Words in [b]ID: #{id}[/b] match slug words: [b]{Links.slug}[/b]. Fix\n'
                    links += error
        self.links_view.text = links

    #this displays the link errors in the tab
    def show_links(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.links_view)

    #this gets all paragraph errors 
    def get_paragraphs(self, *args):
        too_long_paragraphs = ""
        if Paragraphs.h_tag_errors:
            for i in Paragraphs.h_tag_errors:
                error_message = f'[b]Wrong implementation of <h> tags[/b]. Check issue: \n {i}\n\n'
                
        if Paragraphs.long_paragraphs:
            for i in Paragraphs.long_paragraphs:
                error_message = f'[b]Paragraph too long. {i.paragraph_len} words with target 70[/b]: {i.paragraph}\n'
                too_long_paragraphs+=error_message
            
        if Paragraphs.country_modifier_errors:
            for paragraph in Paragraphs.country_modifier_errors:
                error_message = f'\n[b][color=ff0000]!!![/color] Too many country modifiers[/b] in paragraph: \n {paragraph}'
                too_long_paragraphs += error_message
        self.paragraph_view.text=too_long_paragraphs

    #this shows all paragraph errors in the tab, bound to button press
    def show_paragraphs(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.paragraph_view)

    #this gets the bolding error data
    def get_bolding(self, *args):
        if Bolding_Errors.all_bolding_errors:
            error_messages = ''
            for i in Bolding_Errors.all_bolding_errors:
                error = f'[b]{i.bolding_error}[/b]: {i.line_string}\n'
                error_messages+=error
            self.bolding_view.text = error_messages
        else:
            self.bolding_view.text = "No bolding errors found"

    #this shows all bolding errors, bound to button press
    def show_bolding(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.bolding_view)

    #this adds data to the final KW table
    def get_final_table(self, *args): 

        for i in Keyword.complete_data:
            self.keyword_table.add_row(
                (i.keyword, i.exact_match, i.partial_match, i.exact_match + i.partial_match)
            )
        self.parse_button.md_bg_color='red'

    #this shows the KW table, bound to button press
    def show_table(self, *args):
        self.display_section.clear_widgets()
        self.display_section.add_widget(self.keyword_table)

    #this passes the domain name to the Links class
    def pass_site(self, *args):
        link = self.site.text.strip()
        if link[0:8] != 'https://':
            link = 'https://' + link
        Links.domain = link

    #below is file selection fuction
    def select_kw_list_file(self, *args):
        #from plyer import filechooser
        file = filechooser.open_file(on_selection=self.selected_kw_list, filters=['*.txt'])
        if file:
            self.kw_list.text = file[0]
        else:
            pass

    def selected_kw_list(self, selection):
        return selection

    #below is file selection fuction
    def select_docx_file(self, *args):
        file = filechooser.open_file(on_selection=self.selected_docx_file, filters=['*.docx'])
        if file:
            self.docx.text=file[0]
        else:
            pass

    def selected_docx_file(self, selection):
        return selection
    
    

    def main_parse(self, *args): #main function, runs required functions in order
        try:
            keyword_list = self.get_keywords_from_list()
            KW = self.pass_keywords_to_class(keyword_list)
            new_doc = self.word_document()
            formatted_doc = self.format(new_doc)
            self.check_heading(new_doc)
            self.catch_spelling_errors(formatted_doc,database)
            self.exact_match_parse(formatted_doc, KW)
            self.partial_match_parse(formatted_doc, keyword_list, KW)
            Keyword.complete_data = KW
            del KW
            os.remove('formatted_temp.txt')
            os.remove('temp.txt')            
            self.get_final_table()
            self.get_bolding()
            self.get_paragraphs()            
            self.get_links()
            self.get_seo()
            self.get_spelling_errors()
            self.get_spelling_references()
            self.parse_button.md_bg_color='blue'

        #triggers when no files are passed in either docx or txt path
        except FileNotFoundError: 
            self.file_not_found_error = MDDialog(text="File not found")
            self.file_not_found_error.open()

    def get_keywords_from_list(self): #txt file which passed and every line is extracted as a keyword
        list_file = self.kw_list.text
        keyword_list = []
        with open(list_file) as file:
            file = file.readlines()
            for line in file:
                keyword_list.append(line.strip())
        return keyword_list


    def pass_keywords_to_class(self, keyword_list): #creates Keyword class instances and the variable that houses them
        KW = []
        for keyword in keyword_list:
            KW.append(Keyword(keyword))
        return KW

    #this interacts with the .docx file to make its contents into .txt file
    def word_document(self):
        new_doc = open('temp.txt', 'w', errors='ignore')
        document = Document(self.docx.text)

        meta = ""
        title_tag = ""
        for number, table in enumerate(document.tables): #this fucking thing parses all the tables for slug, SEO title and meta description, shit on efficiency
            for num, row in enumerate(table.rows): #if you ever add a tool for parsing alt text and image names, add it here
                for n, cell in enumerate(row.cells):
                    if cell.text.lower() == "meta:":
                        meta = document.tables[number].rows[num].cells[n+1].text.strip()
                    if cell.text.lower() == "title tag:":
                        title_tag = document.tables[number].rows[num].cells[n+1].text.strip()
                    if cell.text.lower() == "slug:":
                        Links.slug = document.tables[number].rows[num].cells[n+1].text.strip()
        SEO_info = SEO(meta, title_tag)
        SEO.info = SEO_info
        del meta, title_tag, SEO_info

        for i in document.paragraphs:
            new_doc.write(i.text + '\n')
        return 'temp.txt'

    def format(self, new_doc): #does a bunch of checks, then formats the file to remove html, punctuation, and other unwanted shit
        with open(new_doc, 'r') as original_file, open('formatted_temp.txt', 'w') as formatted_file:
            long_paragraphs = []
            bolding_errors=[]
            links = []
            class_links = []
            link_errors = []
            h_tag_errors = []
            
            for line in original_file.readlines():
               
                #link tracking tool
                if re.findall(r'id="([^"]+)"', line): #this extracts html id's
                    Links.document_ids.extend(re.findall(r'id="([^"]+)"', line))
                if re.findall(r'<a[^>]*>.*?</a>',line): #this finds all <a> tags
                    if len(re.findall(r'<a[^>]*>.*?</a>',line))>1: #if more than 1 link in paragraph, paragraph sent to errors
                        link_errors.append(line)
                    links.extend(re.findall(r'<a[^>]*>.*?</a>', line))
                if len(re.findall("<a.*?>", line)) != len(re.findall("</a>", line)): #find unmatched <a> tags
                    link_errors.append(line)
                

                #paragraph length check starts here
                if len(line.split(" ")) > 70: 
                    long_paragraphs.append(Paragraphs(line, len(line.split(" "))))
                    
                #heading correctness check
                if line.strip().startswith("<h1") and not line.strip().endswith("</h1>"):
                    h_tag_errors.append(line)
                elif line.strip().startswith("<h2") and not line.strip().endswith("</h2>"):
                    h_tag_errors.append(line)
                elif line.strip().startswith("<h3") and not line.strip().endswith("</h3>"):
                    h_tag_errors.append(line)
                elif line.strip().startswith("<h4") and not line.strip().endswith("</h4>"):
                    h_tag_errors.append(line)
                else:
                    pass

                #bolding check starts here
                if line == "\n" or line.startswith("<h") or line.startswith("<li"): 
                    continue
                elif len(line.split(" "))<25:
                    continue
                elif line.count("<strong>") != line.count("</strong>"):
                    bolding_errors.append(Bolding_Errors(line, "Unmatched <strong> tags"))
                else:
                    pattern = r'<strong>.*?</strong>'
                    found_patterns = reg.findall(pattern, line, overlapped=True)
                    if len(found_patterns)<1:
                        bolding_errors.append(Bolding_Errors(line, "No boldings"))
                    elif len(found_patterns)>1:
                        bolding_errors.append(Bolding_Errors(line, "Too many boldings"))
                    else:
                        pass #bolding check ends here
                
                clean = re.compile('<.*?>')
                line = re.sub(clean,'', line)
                line = re.sub("[,.?:']",'', line)
                line = re.sub("\[[0-9]\]", '', line)
                formatted_file.write(line)

            for link in links:
                class_links.append(Links(link))
            Links.all_links = class_links
            Links.link_errors = link_errors
            Bolding_Errors.all_bolding_errors = bolding_errors
            Paragraphs.long_paragraphs = long_paragraphs
            Paragraphs.h_tag_errors = h_tag_errors

            del long_paragraphs, bolding_errors, links, class_links, link_errors, h_tag_errors

        return 'formatted_temp.txt'
    
    def check_heading(self, new_doc):
        set_country = []
        errors = []
        with open(new_doc, 'r') as file:
            #this loops extracts the country code or country name from the h1 tag
            for line in file.readlines(): 
                if line.strip().startswith("<h1") and line.strip().endswith("</h1>"):
                    
                    heading = line.strip()
                    set_country = {}
                    for country in Paragraphs.countries:
                        if len(country)<=2: #fuck the americans for forcing me to come up with this bullshit. name your fucking country one thing only
                            if country[0] in heading:
                                set_country = [country[0], country[1]]
                                break
                            elif country[1] in heading:
                                set_country = [country[0], country[1]]
                                break
                        else:
                            if country[0] in heading:
                                set_country = [country[0], country[1], country[2]]     
                                break
                            elif country[1] in heading:
                                set_country = [country[0], country[1], country[2]]
                                break
                            elif country[2] in heading:
                                set_country = [country[0], country[1], country[2]]
                                break  
            
        with open(new_doc, 'r') as file:
            #this loop checks all paragraphs for 2 or more country modifiers
            for line in file.readlines():
                if len(set_country)<3: 
                    if line.count(set_country[0]) + line.count(set_country[1]) > 2:
                        errors.append(line)
                    else:
                        continue
                else:                    
                    if (line.count(set_country[0]) + line.count(set_country[1]) + line.count(set_country[2])) > 2:
                        errors.append(line)
                    else:
                        continue
        Paragraphs.country_modifier_errors = errors
        
    def catch_spelling_errors(self, formatted_file, database): #this is some real degenerate shit, but it works most of the time as intended
        single_word_items = []
        double_word_items = []
        spelling_errors = []
        
        #the usefulness of this method is highly questionable. Right now it hogs resources like a motherfucker. 
        #slows down the parsing process by at least 10-20 secs. I expect it would cause shitty laptops to catch on fire.
        #I would just scrap it if it didn't mean I had to remove the database list which I kinda like
        #maybe add a toggle that would disable it at some point

        for line in database:
            if len(line.split(" ")) == 1:
                single_word_items.append(line.replace("\n", "").strip())
            elif len(line.split(" ")) == 2:
                double_word_items.append(line.replace("\n", "").strip())            
            else:
                continue

        commonly_used_words = ['casino','Casino','sports','Sports','gaming','Gaming','games',
        'Games','gambling','Gambling','bet','Bet','slot','Slot','Vegas']

        with open(formatted_file, 'r') as file:
            
            for casino_string in file.readlines():                
                for item in single_word_items: #single word check
                    for phrase in casino_string.split(" "):
                        
                        current_position = casino_string.split(" ").index(phrase)
                        end_position = casino_string.split(" ").index(phrase) + 1
                        
                        window_phrase = (" ".join(casino_string.split(" ")[current_position:end_position])).strip().replace(".", "")
                        if window_phrase in commonly_used_words or window_phrase in single_word_items:
                            continue
                        else:
                            if difflib.SequenceMatcher(None,item, window_phrase).ratio() > 0.75 and difflib.SequenceMatcher(None, window_phrase, item).ratio() < 1:
                                if Spelling(window_phrase, item, casino_string) in spelling_errors:
                                    pass
                                else:
                                    spelling_errors.append(Spelling(window_phrase, item, casino_string))
                                #print(f"1.Potential mispelling. Consider replacing |{window_phrase}| with |{item}|")
                            else:
                                continue

                for item in double_word_items: #double word check
                    for phrase in casino_string.split(" "):
                        current_position = casino_string.split(" ").index(phrase)
                        end_position = casino_string.split(" ").index(phrase) + 2
                        window_phrase = (" ".join(casino_string.split(" ")[current_position:end_position])).strip().replace(".", "")
                        if difflib.SequenceMatcher(None, window_phrase, item).ratio() > 0.77 and difflib.SequenceMatcher(None, window_phrase, item).ratio() < 1:
                            if Spelling(window_phrase, item, casino_string) in spelling_errors:
                                    pass
                            else:
                                spelling_errors.append(Spelling(window_phrase, item, casino_string))
                            #print(f"2.Potential mispelling. Consider replacing |{window_phrase}| with |{item}|")
                        else:
                            continue
                

        Spelling.spelling_errors = spelling_errors
        del spelling_errors, single_word_items, double_word_items
        
    #does what the method name says
    def exact_match_parse(self, formatted_file, KW): 

        formatted_document = open(formatted_file).read()
        keyword_processor = KeywordProcessor()
        for i in KW:
            keyword_processor.add_keyword(i.keyword)
        kw_found = keyword_processor.extract_keywords(formatted_document)
        for i in KW:
            i.exact_match += kw_found.count(i.keyword)

    #does what the method name says. regex works mostly as intended.
    def partial_match_parse(self, formatted_file, keyword_list, KW):
        kw_patterns = []
        for keyword in keyword_list:
            separated_kw = keyword.split(" ")
            kw_patterns.append('(?: |(?:(?: \w* )(?:\w* )?))'.join(separated_kw))
        with open(formatted_file, 'r') as file:
            file = file.read()
            for number, pattern in enumerate(kw_patterns):
                matches = re.findall(pattern, file, re.IGNORECASE)
                matches = [x.lower() for x in matches]
                for match in matches:
                    if match in keyword_list:
                        continue
                    else:
                        KW[number].partial_match +=1

if __name__ == "__main__":
    KeywordParser().run()